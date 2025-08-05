import os

# Fix for Hugging Face Spaces: set Streamlit config to /tmp
os.makedirs("/tmp/.streamlit", exist_ok=True)

with open("/tmp/.streamlit/config.toml", "w") as f:
    f.write("""
[general]
email = ""
""")

os.environ["STREAMLIT_CONFIG_DIR"] = "/tmp/.streamlit"

import streamlit as st
import torch
from transformers import (
    AutoTokenizer, 
    AutoModelForQuestionAnswering,
    pipeline
)
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
import PyPDF2
import docx
import tempfile
import os
import re
import time
from datetime import datetime
from typing import List, Dict, Any
import warnings
warnings.filterwarnings("ignore")

# Set page config
st.set_page_config(
    page_title="Defense Tender LLM Analyzer",
    page_icon="🤖",
    layout="wide"
)

class DefenseTenderLLM:
    def __init__(self):
        """Initialize the Defense Tender LLM System for Hugging Face Spaces"""
        self.documents = []
        self.embeddings = None
        self.faiss_index = None
        self.models_loaded = False
        
    @st.cache_resource
    def load_models(_self):
        """Load models with Hugging Face optimization"""
        try:
            with st.spinner("🤖 Loading AI models... (First time may take 2-3 minutes)"):
                models = {}
                
                progress = st.progress(0)
                status = st.empty()
                
                # 1. Question Answering Model
                status.text("Loading Question Answering model...")
                qa_model_name = "distilbert-base-cased-distilled-squad"
                models['qa_tokenizer'] = AutoTokenizer.from_pretrained(qa_model_name)
                models['qa_model'] = AutoModelForQuestionAnswering.from_pretrained(qa_model_name)
                models['qa_pipeline'] = pipeline(
                    "question-answering",
                    model=models['qa_model'],
                    tokenizer=models['qa_tokenizer']
                )
                progress.progress(25)
                
                # 2. Sentence Transformer for Embeddings
                status.text("Loading Semantic Search model...")
                models['embedding_model'] = SentenceTransformer('all-MiniLM-L6-v2')
                progress.progress(50)
                
                # 3. Text Generation (optional)
                status.text("Loading Text Generation model...")
                try:
                    models['text_generator'] = pipeline(
                        "text-generation",
                        model="distilgpt2",
                        max_length=150,
                        temperature=0.7,
                        pad_token_id=50256
                    )
                except Exception as e:
                    st.warning(f"Text generation unavailable: {e}")
                    models['text_generator'] = None
                progress.progress(75)
                
                # 4. Classification Model (optional)
                status.text("Loading Classification model...")
                try:
                    models['classifier'] = pipeline(
                        "zero-shot-classification",
                        model="facebook/bart-large-mnli"
                    )
                except Exception as e:
                    st.warning(f"Classification unavailable: {e}")
                    models['classifier'] = None
                
                progress.progress(100)
                progress.empty()
                status.empty()
                
                st.success("✅ All models loaded successfully!")
                return models
                
        except Exception as e:
            st.error(f"❌ Error loading models: {e}")
            return None
    
    def extract_text_from_pdf(self, uploaded_file) -> str:
        """Extract text from uploaded PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages[:20]):  # Limit pages
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except:
                    continue
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, uploaded_file) -> str:
        """Extract text from uploaded DOCX file"""
        try:
            doc = docx.Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, uploaded_file) -> str:
        """Extract text from uploaded TXT file"""
        try:
            # Try different encodings
            text_content = uploaded_file.read()
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    if isinstance(text_content, bytes):
                        return text_content.decode(encoding)
                    else:
                        return str(text_content)
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            st.error(f"Error reading TXT: {e}")
            return ""
    
    def intelligent_chunking(self, text: str, tokenizer, chunk_size: int = 400) -> List[str]:
        """Chunk text intelligently for transformers"""
        if not text or not text.strip():
            return []
        
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            test_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            try:
                tokens = tokenizer.encode(test_chunk, add_special_tokens=True)
                if len(tokens) <= chunk_size:
                    current_chunk = test_chunk
                else:
                    if current_chunk.strip():
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence
            except:
                # Fallback to character count
                if len(test_chunk) <= chunk_size * 4:
                    current_chunk = test_chunk
                else:
                    if current_chunk.strip():
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if len(chunk.split()) > 10]
    
    def process_documents(self, uploaded_files, models) -> bool:
        """Process uploaded documents"""
        try:
            self.documents = []
            
            if not uploaded_files:
                st.error("No files uploaded!")
                return False
            
            total_files = len(uploaded_files)
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for idx, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"🤖 Processing: {uploaded_file.name} ({idx+1}/{total_files})")
                
                # Reset file pointer
                uploaded_file.seek(0)
                
                try:
                    # Extract text based on file type
                    text = ""
                    if uploaded_file.name.lower().endswith('.pdf'):
                        text = self.extract_text_from_pdf(uploaded_file)
                    elif uploaded_file.name.lower().endswith('.docx'):
                        text = self.extract_text_from_docx(uploaded_file)
                    elif uploaded_file.name.lower().endswith('.txt'):
                        text = self.extract_text_from_txt(uploaded_file)
                    
                    if text and text.strip():
                        # Chunk the text
                        chunks = self.intelligent_chunking(text, models['qa_tokenizer'])
                        
                        if chunks:
                            for chunk_idx, chunk in enumerate(chunks):
                                self.documents.append({
                                    'text': chunk,
                                    'source': uploaded_file.name,
                                    'chunk_id': len(self.documents),
                                    'chunk_index': chunk_idx,
                                    'total_chunks': len(chunks)
                                })
                            
                            st.success(f"✅ {uploaded_file.name}: {len(chunks)} chunks processed")
                        else:
                            st.warning(f"⚠️ No meaningful content in {uploaded_file.name}")
                    else:
                        st.warning(f"⚠️ Could not extract text from {uploaded_file.name}")
                
                except Exception as e:
                    st.error(f"❌ Error processing {uploaded_file.name}: {e}")
                    continue
                
                progress_bar.progress((idx + 1) / total_files)
            
            if not self.documents:
                st.error("❌ No text content found in any files!")
                return False
            
            # Create embeddings
            self.create_semantic_embeddings(models['embedding_model'])
            
            progress_bar.progress(1.0)
            status_text.text(f"✅ Processed {len(self.documents)} chunks from {total_files} files")
            time.sleep(1)
            progress_bar.empty()
            status_text.empty()
            
            return True
            
        except Exception as e:
            st.error(f"❌ Error in document processing: {e}")
            return False
    
    def create_semantic_embeddings(self, embedding_model):
        """Create semantic embeddings using sentence transformer"""
        if not self.documents:
            return
        
        texts = [doc['text'] for doc in self.documents]
        
        with st.spinner("🧠 Creating semantic embeddings..."):
            # Generate embeddings
            self.embeddings = embedding_model.encode(
                texts,
                convert_to_tensor=False,
                show_progress_bar=False,
                batch_size=32
            )
            
            # Create FAISS index
            dimension = self.embeddings.shape[1]
            self.faiss_index = faiss.IndexFlatIP(dimension)
            
            # Normalize and add embeddings
            faiss.normalize_L2(self.embeddings.astype('float32'))
            self.faiss_index.add(self.embeddings.astype('float32'))
        
        st.success(f"✅ Created {dimension}D embeddings for {len(self.documents)} chunks")
    
    def semantic_search(self, query: str, embedding_model, top_k: int = 5) -> List[Dict]:
        """Perform semantic search"""
        if self.faiss_index is None or not query.strip():
            return []
        
        # Encode query
        query_embedding = embedding_model.encode([query.strip()])
        faiss.normalize_L2(query_embedding.astype('float32'))
        
        # Search
        scores, indices = self.faiss_index.search(query_embedding.astype('float32'), top_k)
        
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(self.documents) and score > 0.1:
                results.append({
                    'text': self.documents[idx]['text'],
                    'source': self.documents[idx]['source'],
                    'score': float(score),
                    'chunk_id': self.documents[idx]['chunk_id']
                })
        
        return results
    
    def answer_question(self, question: str, models) -> Dict[str, Any]:
        """Main question answering pipeline"""
        if not self.documents:
            return {
                'answer': "No documents processed. Please upload documents first.",
                'confidence': 0.0,
                'sources': []
            }
        
        if not question or not question.strip():
            return {
                'answer': "Please enter a valid question.",
                'confidence': 0.0,
                'sources': []
            }
        
        try:
            start_time = time.time()
            
            # Semantic search
            relevant_chunks = self.semantic_search(question.strip(), models['embedding_model'], top_k=5)
            
            if not relevant_chunks:
                return {
                    'answer': "No relevant information found. Try rephrasing your question.",
                    'confidence': 0.0,
                    'sources': []
                }
            
            # Prepare context for QA
            context_texts = []
            total_length = 0
            max_context_length = 400
            
            for chunk in relevant_chunks:
                chunk_tokens = len(models['qa_tokenizer'].encode(chunk['text']))
                if total_length + chunk_tokens <= max_context_length:
                    context_texts.append(chunk['text'])
                    total_length += chunk_tokens
                else:
                    break
            
            context = " ".join(context_texts)
            
            # Get answer using QA model
            result = models['qa_pipeline'](
                question=question,
                context=context
            )
            
            processing_time = time.time() - start_time
            
            return {
                'answer': result['answer'],
                'confidence': result['score'],
                'sources': list(set([chunk['source'] for chunk in relevant_chunks[:3]])),
                'relevant_chunks': relevant_chunks[:3],
                'processing_time': processing_time
            }
            
        except Exception as e:
            return {
                'answer': f"Error processing question: {str(e)}",
                'confidence': 0.0,
                'sources': []
            }

def main():
    """Main Streamlit application for Hugging Face Spaces"""
    st.title("🤖 Defense Tender LLM Analyzer")
    st.markdown("**AI-Powered Document Analysis - Hugging Face Spaces Edition**")
    
    # Model info
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("🧠 AI Models", "4 Transformers")
    with col2:
        st.metric("🔢 Parameters", "570M+")
    with col3:
        st.metric("🏗️ Architecture", "BERT + GPT")
    with col4:
        st.metric("🔍 Search", "Semantic")
    
    # Initialize system
    if 'llm_analyzer' not in st.session_state:
        st.session_state.llm_analyzer = DefenseTenderLLM()
        st.session_state.documents_processed = False
        st.session_state.models = None
    
    # Load models
    if st.session_state.models is None:
        st.session_state.models = st.session_state.llm_analyzer.load_models()
        if st.session_state.models is None:
            st.error("❌ Failed to load models. Please refresh the page.")
            st.stop()
    
    # Sidebar for document upload
    with st.sidebar:
        st.header("📁 Document Upload")
        
        uploaded_files = st.file_uploader(
            "Upload tender documents",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Supported: PDF, DOCX, TXT files"
        )
        
        if uploaded_files:
            st.info(f"📄 {len(uploaded_files)} file(s) selected")
            
            if st.button("🤖 Process with AI", type="primary"):
                start_time = time.time()
                success = st.session_state.llm_analyzer.process_documents(
                    uploaded_files, st.session_state.models
                )
                processing_time = time.time() - start_time
                
                st.session_state.documents_processed = success
                
                if success:
                    st.success(f"✅ Processing completed in {processing_time:.1f}s!")
                    st.balloons()
                else:
                    st.error("❌ Processing failed.")
        
        # Statistics
        if st.session_state.documents_processed:
            st.subheader("📊 Statistics")
            st.metric("Document Chunks", len(st.session_state.llm_analyzer.documents))
            if st.session_state.llm_analyzer.embeddings is not None:
                st.metric("Embedding Dims", st.session_state.llm_analyzer.embeddings.shape[1])
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("🤖 AI Question Answering")
        
        if not st.session_state.documents_processed:
            st.info("👈 Please upload and process documents first.")
            
            with st.expander("💡 About this AI System"):
                st.markdown("""
                **Large Language Models in use:**
                
                🧠 **Question Answering**: DistilBERT (66M parameters)
                - Transformer encoder architecture
                - Fine-tuned for extractive QA
                
                🔍 **Semantic Search**: MiniLM (22M parameters)  
                - Sentence transformer model
                - 384-dimensional embeddings
                
                📝 **Text Generation**: DistilGPT2 (82M parameters)
                - Transformer decoder architecture
                - Autoregressive text generation
                
                🏷️ **Classification**: BART-Large (400M parameters)
                - Encoder-decoder transformer
                - Zero-shot classification
                
                **Total: 570M+ parameters across all models**
                """)
        else:
            # Quick questions
            st.subheader("🎯 Quick Questions")
            example_questions = [
                "What are the eligibility requirements?",
                "When is the submission deadline?",
                "What are the technical specifications?",
                "What is the contract value?",
                "How will proposals be evaluated?",
                "What documents are required?",
                "What are the delivery timelines?",
                "Are there security clearance requirements?"
            ]
            
            cols = st.columns(2)
            for i, question in enumerate(example_questions):
                with cols[i % 2]:
                    if st.button(question, key=f"q_{i}"):
                        st.session_state.selected_question = question
                        st.rerun()
            
            # Custom question
            st.subheader("✍️ Ask Your Question")
            default_question = st.session_state.get('selected_question', '')
            custom_question = st.text_area(
                "Your Question:",
                value=default_question,
                height=100,
                placeholder="Ask anything about the documents...",
                key="question_input"
            )
            
            if st.button("🤖 Get AI Answer", type="primary") and custom_question.strip():
                start_time = time.time()
                
                with st.spinner("🤖 AI is analyzing your question..."):
                    result = st.session_state.llm_analyzer.answer_question(
                        custom_question.strip(), st.session_state.models
                    )
                
                answer_time = time.time() - start_time
                
                # Display answer
                st.subheader("🤖 AI Answer")
                st.write(result['answer'])
                
                # Metrics
                col_conf, col_time, col_sources = st.columns(3)
                
                with col_conf:
                    confidence_color = "green" if result['confidence'] > 0.7 else "orange" if result['confidence'] > 0.4 else "red"
                    st.markdown(f"**Confidence:** :{confidence_color}[{result['confidence']:.1%}]")
                
                with col_time:
                    st.markdown(f"**Response Time:** {answer_time:.2f}s")
                
                with col_sources:
                    st.markdown(f"**Sources:** {len(result.get('sources', []))}")
                
                # Sources
                if result.get('sources'):
                    st.subheader("📚 Source Documents")
                    for source in result['sources']:
                        st.info(f"📄 {source}")
                
                # Relevant chunks
                if 'relevant_chunks' in result and result['relevant_chunks']:
                    with st.expander("🔍 Relevant Text Sections"):
                        for i, chunk in enumerate(result['relevant_chunks'], 1):
                            st.markdown(f"**Section {i}** (Similarity: {chunk['score']:.3f})")
                            st.markdown(f"*Source: {chunk['source']}*")
                            st.markdown(chunk['text'][:400] + "..." if len(chunk['text']) > 400 else chunk['text'])
                            st.markdown("---")
                
                # Clear selected question
                if 'selected_question' in st.session_state:
                    del st.session_state.selected_question
    
    with col2:
        st.header("🚀 System Status")
        
        if st.session_state.models:
            st.success("🟢 AI Models Loaded")
            if st.session_state.documents_processed:
                st.success(f"🟢 {len(st.session_state.llm_analyzer.documents)} chunks ready")
                if st.session_state.llm_analyzer.faiss_index:
                    st.success("🟢 Semantic search active")
            else:
                st.warning("🟡 No documents processed")
        else:
            st.error("🔴 AI models not loaded")
        
        st.subheader("💡 How to Use")
        st.markdown("""
        1. **Upload** your tender documents (PDF, DOCX, TXT)
        2. **Process** them with AI models
        3. **Ask** questions in natural language
        4. **Get** intelligent answers with sources
        
        The AI uses semantic understanding, not just keyword matching!
        """)
        
        if st.session_state.documents_processed:
            st.subheader("🧹 Actions")
            if st.button("Clear Cache"):
                # Clear cache
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                st.success("Cache cleared!")

if __name__ == "__main__":
    main()
